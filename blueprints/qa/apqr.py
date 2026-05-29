"""
KD-IRIS — APQR(연간제품품질평가) 초안 자동작성 (Phase M-3)

품목(item_seq) + 연도 → 규제 이벤트 + 마스터 정보를 APQR 13섹션 초안으로 자동 생성.
규제 섹션(1~7)은 KD-IRIS 데이터로 자동 채움, SAP/EDMS 섹션(8~12)은 'pending' 표시,
종합(13)은 1~7 요약 서술(룰베이스 — MISO drug_consult_agent 로 교체 가능).

문서 출력: export_apqr_docx() → python-docx 로 Word(.docx) 생성.
"""
import datetime
import io
import logging
from typing import Dict, List, Optional

from . import db as qadb
from .integrations import fetch_sap_batches, fetch_edms_documents

logger = logging.getLogger(__name__)


# ────────────────────────────────────────────────────────────────────────────
def _year_events(item_code: str, year: int) -> List[Dict]:
    """품목(사내코드)의 규제 이벤트 — event.matched_item_codes 기반(lookup 재사용) + 연도 필터."""
    y = str(year)
    try:
        from . import lookup as qa_lookup
        rows = qa_lookup.events_for_item(item_code, limit=200)
    except Exception:
        rows = []
    # event_date(YYYYMMDD) 가 해당 연도이거나, 날짜 미상이면 보수적으로 포함
    return [e for e in rows if (e.get("event_date") or "").startswith(y) or not e.get("event_date")]


def _summary_text(p: Dict, recalls, discs, letters, reviews, ingredients) -> str:
    """종합 평가 서술 (룰베이스 초안). MISO 연동 시 이 함수를 교체."""
    name = p.get("item_name") or p.get("item_code")
    parts = [f"본 평가는 {name}의 당해 연도 규제·품질 관련 외부 이벤트를 종합한 것이다."]
    if recalls:
        parts.append(f"회수·판매중지 {len(recalls)}건이 확인되어 해당 배치 영향 평가 및 시정조치 이행 여부 점검이 필요하다.")
    else:
        parts.append("당해 연도 회수·판매중지 이력은 없다.")
    if discs:
        parts.append(f"행정처분 {len(discs)}건이 확인되어 위반사항의 자사 적용 여부 검토가 요구된다.")
    else:
        parts.append("행정처분 이력은 확인되지 않았다.")
    if letters:
        parts.append(f"안전성서한 {len(letters)}건과 관련하여 라벨·DUR 반영 여부를 확인한다.")
    if reviews:
        parts.append(f"재심사·재평가 관련 일정 {len(reviews)}건을 모니터링 중이다.")
    elif p.get("reexam_date"):
        parts.append(f"재심사 기한일은 {p.get('reexam_date')}이다.")
    if not (recalls or discs or letters):
        parts.append("종합적으로 당해 연도 중대한 규제 이벤트는 없어 품질 상태는 안정적인 것으로 판단된다.")
    parts.append("제조·시험 결과(SAP), 일탈·CAPA(EDMS), 안정성 시험(LIMS) 데이터는 시스템 연동 후 보완 평가가 필요하다.")
    return " ".join(parts)


def build_apqr(item_code: str, year: int) -> Optional[Dict]:
    """품목(사내 품목코드) + 연도 → APQR 초안. 규제 섹션 자동(master_item+event), SAP/EDMS 'pending'."""
    p = qadb.query_one("SELECT * FROM master_item WHERE item_code=?", (item_code,))
    if not p:
        return None
    ingredients = qadb.query("SELECT * FROM master_ingredient WHERE item_code=?", (item_code,))
    yev = _year_events(item_code, year)

    def by_type(*types):
        return [e for e in yev if e.get("event_type") in types]

    recalls = by_type("recall")
    discs = by_type("disciplinary")
    letters = by_type("safety_letter")
    reviews = by_type("review_due", "reeval_done")
    suppliers = by_type("supplier_closure")

    sap = fetch_sap_batches(item_code, year)
    edms = fetch_edms_documents(item_code, ["std", "deviation", "capa"])

    today = datetime.date.today().isoformat()
    meta = {
        "item_code": item_code,
        "item_seq": p.get("item_seq") or "—",
        "product_name": p.get("item_name"),
        "permit_no": p.get("permit_no"),
        "permit_holder": p.get("entp_name"),
        "etc_otc": p.get("category"),
        "year": year,
        "period": f"{year}-01-01 ~ {year}-12-31",
        "created_at": today,
        "reexam_date": p.get("reexam_date"),
    }

    sections = [
        {"no": 1, "title": "제품 개요 및 허가사항", "source": "MASTER", "auto": True,
         "rows": [
             ("제품명", p.get("item_name")), ("사내 품목코드", p.get("item_code")),
             ("품목기준코드", p.get("item_seq")), ("허가번호", p.get("permit_no")),
             ("허가일자", p.get("permit_date")), ("허가권자", p.get("entp_name")),
             ("의약품구분", p.get("category")), ("분류", p.get("sub_category")),
             ("성상", p.get("chart")), ("보관조건", p.get("storage_method")),
             ("사용기간", (str(p.get("shelf_life_mo")) + "개월") if p.get("shelf_life_mo") else None),
             ("제조단위", p.get("batch_size")), ("포장규격", p.get("pack_spec")),
             ("보험코드(EDI)", p.get("edi_code")), ("ATC", p.get("atc_code")),
         ]},
        {"no": 2, "title": "허가 변경이력", "source": "API:140", "auto": True,
         "text": f"최초 허가일 {p.get('permit_date') or '미상'}. 당해 연도 허가 변경 이벤트 {len(by_type('permit_change'))}건.",
         "events": by_type("permit_change")},
        {"no": 3, "title": "회수·판매중지 현황", "source": "API:539", "auto": True,
         "summary": f"{len(recalls)}건" if recalls else "해당 없음 (이상 없음)", "events": recalls},
        {"no": 4, "title": "행정처분 현황", "source": "API:564", "auto": True,
         "summary": f"{len(discs)}건" if discs else "해당 없음", "events": discs},
        {"no": 5, "title": "안전성서한·DUR 변경", "source": "API:547/531", "auto": True,
         "summary": f"안전성서한 {len(letters)}건" if letters else "해당 없음", "events": letters},
        {"no": 6, "title": "재심사·재평가 일정", "source": "API:554/556", "auto": True,
         "text": f"재심사 기한일: {p.get('reexam_date') or '해당 없음'}. 관련 이벤트 {len(reviews)}건.",
         "events": reviews},
        {"no": 7, "title": "원료·공급처 변경/평가", "source": "MASTER+API:483", "auto": True,
         "rows": [(i.get("ingr_name"), f"{i.get('manufacturer') or '제조처 미상'} / {i.get('supplier') or '공급처 미상'}")
                  for i in ingredients] or [("원료 정보", "보강 데이터 없음")],
         "events": suppliers},
        {"no": 8, "title": "제조번호 및 제조/시험 결과", "source": "SAP", "auto": False,
         "status": "연동 대기", "rows": sap},
        {"no": 9, "title": "OOS/OOT 현황", "source": "LIMS", "auto": False, "status": "연동 대기"},
        {"no": 10, "title": "안정성 시험 결과", "source": "LIMS", "auto": False, "status": "연동 대기"},
        {"no": 11, "title": "변경관리·일탈·CAPA", "source": "EDMS/QMS", "auto": False,
         "status": "연동 대기", "rows": edms},
        {"no": 12, "title": "클레임·반품", "source": "SAP/QMS", "auto": False, "status": "연동 대기"},
        {"no": 13, "title": "종합 평가 및 조치사항", "source": "AI", "auto": True,
         "text": _summary_text(p, recalls, discs, letters, reviews, ingredients)},
    ]
    return {"meta": meta, "sections": sections}


# ────────────────────────────────────────────────────────────────────────────
# Word(.docx) 출력 — python-docx
# ────────────────────────────────────────────────────────────────────────────
def export_apqr_docx(item_code: str, year: int) -> Optional[io.BytesIO]:
    """APQR 초안을 Word(.docx) BytesIO 로 생성. 데이터 없으면 None."""
    data = build_apqr(item_code, year)
    if not data:
        return None
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    m = data["meta"]
    doc.add_heading(f"연간제품품질평가(APQR) — {m['product_name']}", level=0)
    doc.add_paragraph(f"평가기간: {m['period']}    작성일: {m['created_at']}")
    doc.add_paragraph(f"품목기준코드: {m['item_seq']}    허가번호: {m.get('permit_no') or '-'}    "
                      f"허가권자: {m.get('permit_holder') or '-'}")

    for s in data["sections"]:
        doc.add_heading(f"{s['no']}. {s['title']}", level=1)
        if not s.get("auto"):
            doc.add_paragraph(f"[{s.get('status', '연동 대기')}] — {s['source']} 시스템 연동 후 자동 작성")
            rows = s.get("rows") or []
            if rows:
                _docx_rows_table(doc, rows)
            continue
        if s.get("text"):
            doc.add_paragraph(s["text"])
        if s.get("summary"):
            doc.add_paragraph(f"요약: {s['summary']}")
        rows = s.get("rows")
        if rows:
            _docx_rows_table(doc, rows)
        events = s.get("events")
        if events:
            t = doc.add_table(rows=1, cols=3)
            t.style = "Light Grid Accent 1"
            h = t.rows[0].cells
            h[0].text, h[1].text, h[2].text = "일자", "제목", "심각도"
            for e in events:
                c = t.add_row().cells
                c[0].text = str(e.get("event_date") or "-")
                c[1].text = str(e.get("title") or "")
                c[2].text = str(e.get("severity") or "")

    doc.add_paragraph("")
    note = doc.add_paragraph("※ 본 문서는 KD-IRIS가 규제 공공데이터로 자동 생성한 초안입니다. "
                            "제조·시험·일탈 데이터는 SAP/EDMS/LIMS 연동 후 보완됩니다.")
    note.runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _docx_rows_table(doc, rows):
    """(라벨, 값) 튜플 리스트 → 2열 표."""
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light List Accent 1"
    for label, val in rows:
        c = t.add_row().cells
        c[0].text = str(label or "")
        c[1].text = str(val if val is not None else "-")
